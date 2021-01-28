import requests
import shutil
import json
#from types import SimpleNamespace
import docx

def DownloadPic(image_url):
    filename = image_url.split("/")[-1]
    image = requests.get(image_url, stream = True)
    if image.status_code == 200:
        image.raw.decode_content = True
        with open(filename,'wb') as f:
            shutil.copyfileobj(image.raw, f)

def GetData():
    print("Hello GetData\n")
    url="https://reqres.in/api/users?page=2"
    #url="https://reqres.in/api/users/2"
    response=requests.get(url)
    json_all=json.loads(response.text)
    #json_all=json.loads(response.text, object_hook=lambda d: SimpleNamespace(**d))
    json_data=json_all["data"]
    for data in json_data:
        print(data["avatar"])
        DownloadPic(data["avatar"])
    #for user in json_data:
    #    print(user)
    #print (json_all["data"])
    return json_data

def CreateDoc(json_data):
    mydoc = docx.Document()
    mydoc.add_paragraph("Report of System Users.")
    table = mydoc.add_table(rows=1,cols=5)
    hd_cells = table.rows[0].cells
    hd_cells[0].text = 'User ID'
    hd_cells[1].text = 'User Email'
    hd_cells[2].text = 'User Last Name'
    hd_cells[3].text = 'User First Name'
    hd_cells[4].text = 'User Avatar'
    for data in json_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(data["id"])
        row_cells[1].text = str(data["email"])
        row_cells[2].text = str(data["last_name"])
        row_cells[3].text = str(data["first_name"])
 
        filename = data["avatar"].split("/")[-1]
        para = row_cells[4].paragraphs[0]
        run = para.add_run()
        run.add_picture(filename)
    mydoc.save("new_doc.docx")
        

json_data=GetData()
#SavingUsers(json_data)
CreateDoc(json_data)
#print (json_data)
print("Hello world")
